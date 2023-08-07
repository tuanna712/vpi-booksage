import os
import pdfplumber
import streamlit as st
import tiktoken
from docx import Document as docxDoc
from langchain.schema import Document
from langchain.embeddings import CohereEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from underthesea import word_tokenize
from langchain.vectorstores import Qdrant
from qdrant_client import QdrantClient
from dotenv import load_dotenv
load_dotenv()

class DocProcessing():
    def __init__(self, 
                 uploaded_file,
                 vector_path:str='database/user_1/docs_db',
                 chunk_size:int=500,
                 chunk_overlap:int=100,
                 collection_name:str=None,
                 chunks_plot:bool=False,
                 vectordb:str='non-existed',
                 vmethod:str='qdrant',
                 book_lang:str='en',
                 separators:list=['\n\n\n\n','\n\n\n','\n\n', '\n', ' ', ''],
                 ):
        
        ##Self Definition--------------------------
        self.uploaded_file = uploaded_file
        self.vector_path = vector_path
        self.qdrant_url = os.environ['QDRANT_URL']
        self.qdrant_api_key = os.environ['QDRANT_API_KEY']
        self.embeddings = CohereEmbeddings(model="multilingual-22-12", cohere_api_key=os.environ['COHERE_API_KEY'])
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.collection_name = collection_name
        self.chunks_plot = chunks_plot
        self.vectordb = vectordb
        self.vmethod = vmethod
        self.book_lang = book_lang
        self.separators = separators

        ##Processing...----------------------------
    def file_processing(self):
        self.file_input()
        self.file_loading()
        self.text_splitting()

        if self.vmethod=='chroma':
            self.chroma_vectorization()
        else:
            self.qdrant_vectorization()
            
        st.write('Processed Documents to VectorDatabase!')
        
    def file_input(self):
        if self.uploaded_file is not None:
                #Split file extension for detection of loading type
                self.file_extension = os.path.splitext(self.uploaded_file.name)[-1]
                #Load PDF file
                if self.file_extension == ".pdf":
                    self.document = pdfplumber.open(self.uploaded_file)
                    print('Loaded PDF file!\n')
                #Load MS Word file
                if self.file_extension in [".doc", ".docx"]:
                    self.document = docxDoc(self.uploaded_file)
                    print('Loaded MS Word file!\n')
                return self.document
            
    def file_loading(self):
        ##Process-PDF-------------------------------
        if self.file_extension == ".pdf":
            #Convert Pages to Document which interact with Langchain
            self.data = []
            for i in range(len(self.document.pages)):
                page = Document(page_content=self.document.pages[i].extract_text(), 
                                    metadata={'source': self.uploaded_file.name, 
                                            'page': i})
                self.data.append(page)
        ##Process-MS WORD---------------------------
        if self.file_extension in [".doc", ".docx"]:
            self.data = ""
            for p in self.document.paragraphs:
                self.data = self.data + "\n" + p.text
        
    def text_splitting(self):
        #Split to Smaller Texts
        text_splitter = RecursiveCharacterTextSplitter(
                                chunk_size=self.chunk_size,
                                chunk_overlap=self.chunk_overlap,
                                length_function=self.tiktoken_len,
                                separators=self.separators,
                                )
        ##Process-MS WORD---------------------------
        if self.file_extension in [".doc", ".docx"]:
            self.chunks = text_splitter.split_text(self.data)
            if self.book_lang=='vi':
                self.chunks = [word_tokenize(self.chunks[i], format="text") 
                               for i in range(len(self.chunks))]
            print(len(self.data))
            print(len(self.chunks))
        ##Process-PDF-------------------------------
        if self.file_extension in [".pdf"]:
            self.chunks = []
            for i in range(len(self.data)):
                chunk_contents = text_splitter.split_text(self.data[i].page_content)
                if self.book_lang=='vi':
                    chunk_contents = [word_tokenize(chunk_contents[i], format="text") for i in range(len(chunk_contents))]
                for j in range(len(chunk_contents)):
                    chunk = Document(page_content=chunk_contents[j], 
                                    metadata={'source': self.data[i].metadata['source'], 
                                            'page': self.data[i].metadata['page']})
                    self.chunks.append(chunk)
                pass
            print(len(self.data))
            print(len(self.chunks))
            
    def tiktoken_len(self, text):
        tokenizer = tiktoken.get_encoding('cl100k_base')
        tokens = tokenizer.encode(text, disallowed_special=())
        return len(tokens)
    
    def token_counts_plot(self, token_counts):
        import plotly.express as px
        # Create histogram using Plotly
        fig = px.histogram(token_counts, nbins=30)
        # Customize the plot info
        fig.update_layout(
            title="Chunks-Size Histogram",
            xaxis_title="Chunk-Size",
            yaxis_title="Number of chunks",
            bargap=0.1,
            showlegend=False)
        # Display the plot
        st.plotly_chart(fig, use_container_width=True)
        
    def display_chunks_hist(self):
        if self.file_extension in [".pdf"]:
            token_counts = [self.tiktoken_len(chunk.page_content) for chunk in self.chunks]
        else:
            token_counts = [self.tiktoken_len(chunk) for chunk in self.chunks]
        self.token_counts_plot(token_counts)
        
    def total_number_of_chunks(self):
        return len(self.chunks)
    
    def display_chunks(self, n_chunk):
        return self.chunks
                
    def qdrant_vectorization(self):
        qdrant_url = os.environ['QDRANT_URL']
        qdrant_api_key = os.environ['QDRANT_API_KEY']
        client = QdrantClient(url=qdrant_url, api_key=qdrant_api_key)
        client.delete_collection(collection_name=self.collection_name)
        with st.spinner(text='Embedding...'):
            ids = [i for i in range(len(self.chunks))]
            ##Process-PDF-------------------------------
            if self.file_extension in [".pdf"]:
                ###Online-Mode--------------------------
                self.vdatabase = Qdrant.from_documents(self.chunks,
                                                self.embeddings, 
                                                ids = ids,
                                                url=self.qdrant_url, 
                                                prefer_grpc=True, 
                                                api_key=self.qdrant_api_key, 
                                                collection_name=self.collection_name,
                                                )
            ##Process-DOCX-------------------------------
            if self.file_extension in [".doc", ".docx"]:
                ###Online-Mode--------------------------
                self.vdatabase = Qdrant.from_texts(self.chunks,
                                                self.embeddings, 
                                                ids = ids,
                                                url=self.qdrant_url, 
                                                prefer_grpc=True, 
                                                api_key=self.qdrant_api_key, 
                                                collection_name=self.collection_name,
                                                )
        st.info("Qdrant Vectorized Chunks use on-cloud storage\n")
    
    def chroma_vectorization(self):
        # import chromadb
        # from chromadb.utils import embedding_functions
        # # Client to connect to vector db
        # client = chromadb.PersistentClient(path=self.vector_path)
        # # Embedding function
        # cohere_ef  = embedding_functions.CohereEmbeddingFunction(
        #                             api_key=os.environ['COHERE_API_KEY'], 
        #                             model_name="multilingual-22-12")
        # # get a collection or create if it doesn't exist already
        # try:
        #     _collection = client.get_or_create_collection(self.collection_name, 
        #                                               embedding_function=cohere_ef)
        # except ValueError:
        #     st.warning("""Expected collection name that \n(1) contains 3-63 characters, \n(2) starts and ends with an alphanumeric character, \n(3) otherwise contains only alphanumeric characters, underscores or hyphens (-), \n(4) contains no two consecutive periods (..)""")
        #     st.stop()
        # if self.file_extension in [".doc", ".docx"]:
        #     # upsert items. new items will be added, existing items will be updated.
        #     _collection.upsert(
        #                 ids=['id_'+str(i) for i in range(1, len(self.chunks)+1)],
        #                 documents=[self.chunks[i].replace('\n', ' ') 
        #                            for i in range(len(self.chunks))],
        #                 )
        # elif self.file_extension in [".pdf"]:
        #     # define ids, metadatas and documents
        #     ids = ['id_'+str(i) for i in range(1, len(self.chunks)+1)]
        #     metadatas = [{'source': self.chunks[j].metadata['source'],
        #                     'page': self.chunks[j].metadata['page']} for j in range(len(self.chunks))
        #                  ]
        #     documents = [self.chunks[i].page_content.replace('\n', ' ') for i in range(len(self.chunks))]
        #     # upsert items. new items will be added, existing items will be updated.
        #     _collection.upsert(ids=ids,
        #             metadatas=metadatas,
        #             documents=documents,)
        pass